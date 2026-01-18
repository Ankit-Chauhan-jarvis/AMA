import os
import re
import json
import time
import argparse
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional

from groq import Groq
from dotenv import load_dotenv
from docx import Document


LABEL_SET = ["task", "procedural", "social", "action", "counter"]


def docx_to_text(docx_path: str) -> str:
    """
    Extracts readable text from paragraphs + tables.
    This is used as the main "labeling guide" context for the LLM.
    """
    doc = Document(docx_path)
    parts: List[str] = []

    def norm(s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r"[ \t]+", " ", s)
        s = re.sub(r"\n{3,}", "\n\n", s)
        return s.strip()

    for p in doc.paragraphs:
        t = norm(p.text)
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [norm(c.text) for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def try_extract_examples_from_tables(docx_path: str) -> Dict[str, List[str]]:
    """
    Best-effort: tries to extract example sentences per label from tables
    formatted like:
      Row0 col0: includes "(Label: <code>)"
      Row3 col1: examples
    If not found, returns {} (and we rely on raw guide text).
    """
    doc = Document(docx_path)
    label_header_re = re.compile(r"\(Label:\s*([a-zA-Z_]+)\)", re.IGNORECASE)
    examples_by_label: Dict[str, List[str]] = {}

    def clean(s: str) -> str:
        s = (s or "").strip()
        s = re.sub(r"[ \t]+", " ", s)
        return s.strip()

    for table in doc.tables:
        if not table.rows or len(table.rows) < 4:
            continue

        header = clean(table.rows[0].cells[0].text)
        m = label_header_re.search(header)
        if not m:
            continue

        label = m.group(1).strip().lower()
        if label not in LABEL_SET:
            continue

        try:
            ex_raw = clean(table.rows[3].cells[1].text)
        except Exception:
            ex_raw = ""

        if not ex_raw:
            continue

        items: List[str] = []
        for line in ex_raw.splitlines():
            line = line.strip(" \t-•")
            if not line:
                continue
            for piece in line.split("|"):
                piece = piece.strip(" \t-•")
                if piece:
                    items.append(piece)

        if items:
            examples_by_label[label] = items

    return examples_by_label


def build_system_prompt(labeling_guide_text: str, extracted_examples: Dict[str, List[str]]) -> str:
    """
    System prompt provides:
      - label set
      - the full labeling guide extracted from the DOCX
      - optional few-shot examples (if extractable)
    """
    GUIDE_CHAR_CAP = 15000
    guide = labeling_guide_text
    if len(guide) > GUIDE_CHAR_CAP:
        guide = guide[:GUIDE_CHAR_CAP] + "\n\n[NOTE: Labeling guide truncated due to length.]"

    blocks: List[str] = []
    blocks.append(
        "You are labeling meeting transcript sentences. Each sentence is one sense unit.\n"
        f"Choose exactly ONE label from: {', '.join(LABEL_SET)}.\n"
        "Return ONLY valid JSON as instructed (no markdown, no extra text).\n"
        "If the sentence is a short acknowledgement or backchannel (e.g., 'yeah', 'ok', 'right', 'great'), label it as social.\n"
        "Pick the single best label even if uncertain.\n"
    )

    blocks.append("=== LABELING GUIDE (from provided document) ===")
    blocks.append(guide)

    if extracted_examples:
        blocks.append("\n=== FEW-SHOT EXAMPLES (sentence -> label) ===")
        for label in LABEL_SET:
            exs = extracted_examples.get(label, [])
            for ex in exs[:4]: 
                blocks.append(f"- {ex} -> {label}")

    return "\n".join(blocks).strip()


def read_text_file(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def split_into_sentences(transcript: str) -> List[str]:
    """
    Transcript-friendly splitter:
    - respects line breaks (speaker turns)
    - splits on . ! ? followed by whitespace
    - keeps short utterances
    """
    text = transcript.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]

    sent_split_re = re.compile(r"(?<=[.!?])\s+")
    sentences: List[str] = []
    for line in lines:
        parts = sent_split_re.split(line)
        for p in parts:
            s = p.strip()
            if s:
                sentences.append(s)

    return sentences


def extract_json_array(text: str) -> List[Dict[str, Any]]:
    start = text.find("[")
    end = text.rfind("]")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Model output did not contain a JSON array.")
    return json.loads(text[start : end + 1])


def label_batch(
    client: Groq,
    model: str,
    system_prompt: str,
    batch_items: List[Tuple[int, str]],
    temperature: float = 0.0,
    max_retries: int = 3,
) -> List[Dict[str, Any]]:
    user_prompt = (
        "Label each sentence with exactly one label from: "
        + ", ".join(LABEL_SET)
        + ".\n"
        "Return ONLY a JSON array.\n"
        'Each element must be: {"id": <int>, "label": "<one_of_labels>"}\n\n'
        "Sentences:\n"
    )

    for sid, sent in batch_items:
        user_prompt += f"{sid}: {sent}\n"

    last_err = None
    for attempt in range(1, max_retries + 1):
        try:
            chat_completion = client.chat.completions.create(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                model=model,
                temperature=temperature,
            )

            out = (chat_completion.choices[0].message.content or "").strip()
            arr = extract_json_array(out)

            expected_ids = {sid for sid, _ in batch_items}
            got_ids = set()
            results: List[Dict[str, Any]] = []

            for obj in arr:
                sid = int(obj["id"])
                lab = str(obj["label"]).strip().lower()
                if lab not in LABEL_SET:
                    raise ValueError(f"Invalid label '{lab}' for id={sid}")
                got_ids.add(sid)
                results.append({"id": sid, "label": lab})

            if expected_ids != got_ids:
                missing = sorted(list(expected_ids - got_ids))
                extra = sorted(list(got_ids - expected_ids))
                raise ValueError(f"Batch id mismatch. Missing: {missing}, Extra: {extra}")

            results.sort(key=lambda x: x["id"])
            return results

        except Exception as e:
            last_err = e
            if attempt < max_retries:
                time.sleep(1.5 * attempt)

    raise RuntimeError(f"Failed labeling batch after retries. Last error: {last_err}")


def ensure_exists(p: str, what: str) -> str:
    path = Path(p)
    if not path.exists():
        raise FileNotFoundError(
            f"{what} not found at: {p}\n"
            f"Resolved path: {str(path.resolve())}\n"
            "Fix: pass the correct local Windows path (not /mnt/data/...)."
        )
    return str(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Sentence-level act4teams-short labeling via Groq + gpt-oss-20b.")
    parser.add_argument("--transcript", required=True, help="Path to Transcript_text.txt")
    parser.add_argument("--labels_docx", required=True, help="Path to act4teams_labels_description_en.docx")
    parser.add_argument("--out", default="labeled_sentences.jsonl", help="Output JSONL path")
    parser.add_argument("--batch_size", type=int, default=20, help="Sentences per LLM call")
    args = parser.parse_args()

    load_dotenv()

    transcript_path = ensure_exists(args.transcript, "Transcript file")
    labels_docx_path = ensure_exists(args.labels_docx, "Labels DOCX file")

    GROQ_API_KEY = os.getenv("GROQ_API_KEY")

    if not GROQ_API_KEY:
        raise ValueError("GROQ_API_KEY not found in .env file.")

    client = Groq(api_key=GROQ_API_KEY)

    model = os.getenv("GROQ_MODEL", "openai/gpt-oss-20b")

    labeling_guide_text = docx_to_text(labels_docx_path)
    extracted_examples = try_extract_examples_from_tables(labels_docx_path)  
    system_prompt = build_system_prompt(labeling_guide_text, extracted_examples)

    transcript_text = read_text_file(transcript_path)
    sentences = split_into_sentences(transcript_text)

    indexed: List[Tuple[int, str]] = [(i + 1, s) for i, s in enumerate(sentences)]
    id_to_label: Dict[int, str] = {}

    for i in range(0, len(indexed), args.batch_size):
        batch = indexed[i : i + args.batch_size]
        labeled = label_batch(
            client=client,
            model=model,
            system_prompt=system_prompt,
            batch_items=batch,
            temperature=0.0,
        )
        for item in labeled:
            id_to_label[item["id"]] = item["label"]

    with open(args.out, "w", encoding="utf-8") as f:
        for sid, sent in indexed:
            rec = {"id": sid, "sentence": sent, "label": id_to_label.get(sid, "")}
            f.write(json.dumps(rec, ensure_ascii=False) + "\n")

    print(f"Wrote {len(indexed)} labeled sentences to: {args.out}")


if __name__ == "__main__":
    main()
